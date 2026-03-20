"""
app/services/extractor.py
──────────────────────────
Extracts raw text from uploaded PDF and DOCX files.
Handles edge-cases: scanned PDFs (fallback), password-protected, corrupt.
"""

from __future__ import annotations
import re
from pathlib import Path
from typing import Optional

from app.core.logger import get_logger

logger = get_logger(__name__)


class TextExtractor:
    """
    Unified text extractor supporting .pdf and .docx formats.
    Returns cleaned UTF-8 text ready for NLP processing.
    """

    # ── Public API ─────────────────────────────────────────────────────────────

    def extract(self, filepath: str | Path) -> str:
        """
        Extract text from a file. Auto-detects format from extension.

        Args:
            filepath: Absolute or relative path to the file.

        Returns:
            Cleaned plain text string.

        Raises:
            ValueError: If the file extension is unsupported.
            RuntimeError: If extraction fails.
        """
        path = Path(filepath)
        ext  = path.suffix.lower()

        logger.debug("Extracting text from: %s", path.name)

        if ext == ".pdf":
            text = self._extract_pdf(path)
        elif ext == ".docx":
            text = self._extract_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        cleaned = self._clean(text)
        logger.debug("Extracted %d characters from %s", len(cleaned), path.name)
        return cleaned

    # ── PDF extraction ─────────────────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> str:
        """Use pdfplumber as primary extractor."""
        try:
            import pdfplumber  # type: ignore

            pages: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)

            if pages:
                return "\n".join(pages)

            logger.warning("pdfplumber returned empty text for %s — trying PyPDF2 fallback", path.name)
            return self._extract_pdf_fallback(path)

        except Exception as exc:
            logger.warning("pdfplumber failed (%s) — falling back", exc)
            return self._extract_pdf_fallback(path)

    def _extract_pdf_fallback(self, path: Path) -> str:
        """Fallback: PyPDF2."""
        try:
            from PyPDF2 import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            pages = [
                page.extract_text() or ""
                for page in reader.pages
            ]
            return "\n".join(pages)
        except Exception as exc:
            raise RuntimeError(f"PDF extraction failed entirely: {exc}") from exc

    # ── DOCX extraction ────────────────────────────────────────────────────────

    def _extract_docx(self, path: Path) -> str:
        """Extract text from .docx using python-docx."""
        try:
            from docx import Document  # type: ignore

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            return "\n".join(paragraphs)
        except Exception as exc:
            raise RuntimeError(f"DOCX extraction failed: {exc}") from exc

    # ── Text cleaning ──────────────────────────────────────────────────────────

    def _clean(self, text: str) -> str:
        """
        Normalise extracted text:
        - Collapse multiple blank lines
        - Fix hyphenated line-breaks
        - Normalise whitespace
        """
        if not text:
            return ""

        # Fix hyphenated line-breaks (e.g. "develop-\nment" → "development")
        text = re.sub(r"-\n", "", text)

        # Collapse 3+ consecutive blank lines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Collapse multiple spaces/tabs to single space
        text = re.sub(r"[ \t]+", " ", text)

        # Strip lines
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(lines).strip()


# ── Singleton ──────────────────────────────────────────────────────────────────
extractor = TextExtractor()
